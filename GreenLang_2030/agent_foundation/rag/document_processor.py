"""
Document Processor for GreenLang RAG System

Handles document parsing, chunking, and metadata extraction for multiple formats:
- PDF documents (PyPDF2, pdfplumber)
- HTML content (BeautifulSoup)
- Office documents (python-docx, openpyxl)
- Structured data (CSV, JSON, XML)

Implements smart chunking strategies with overlap for context preservation.
"""

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import numpy as np

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    HTML = "html"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    TXT = "txt"
    MARKDOWN = "md"


class ChunkingStrategy(Enum):
    """Document chunking strategies"""
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    SLIDING_WINDOW = "sliding_window"
    SENTENCE_BASED = "sentence_based"
    RECURSIVE_CHARACTER = "recursive_character"
    PARAGRAPH_BASED = "paragraph_based"
    TOKEN_BASED = "token_based"


@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    source: str
    doc_type: DocumentType
    created_at: datetime = field(default_factory=datetime.now)
    page_count: int = 0
    word_count: int = 0
    language: str = "en"
    author: Optional[str] = None
    title: Optional[str] = None
    tags: List[str] = field(default_factory=list)
    custom_metadata: Dict[str, Any] = field(default_factory=dict)
    provenance_hash: Optional[str] = None

    def calculate_provenance(self, content: str) -> str:
        """Calculate SHA-256 hash for provenance tracking"""
        hash_input = f"{self.source}{content}{self.created_at.isoformat()}"
        self.provenance_hash = hashlib.sha256(hash_input.encode()).hexdigest()
        return self.provenance_hash


@dataclass
class Document:
    """Document with content and metadata"""
    content: str
    metadata: DocumentMetadata
    doc_id: Optional[str] = None
    chunk_id: Optional[int] = None
    embedding: Optional[np.ndarray] = None
    confidence_score: float = 1.0
    start_char: Optional[int] = None
    end_char: Optional[int] = None

    def __post_init__(self):
        if not self.doc_id:
            # Generate unique document ID
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            source_hash = hashlib.md5(self.metadata.source.encode()).hexdigest()[:8]
            self.doc_id = f"{source_hash}_{content_hash}"


class DocumentParser:
    """Parse documents from various formats"""

    def __init__(self):
        self.parsers = {
            DocumentType.PDF: self._parse_pdf,
            DocumentType.HTML: self._parse_html,
            DocumentType.DOCX: self._parse_docx,
            DocumentType.XLSX: self._parse_xlsx,
            DocumentType.CSV: self._parse_csv,
            DocumentType.JSON: self._parse_json,
            DocumentType.XML: self._parse_xml,
            DocumentType.TXT: self._parse_text,
            DocumentType.MARKDOWN: self._parse_markdown,
        }

    def parse(
        self,
        file_path: Union[str, Path],
        doc_type: Optional[DocumentType] = None,
        extract_metadata: bool = True
    ) -> Tuple[str, DocumentMetadata]:
        """
        Parse document and extract content with metadata

        Args:
            file_path: Path to the document
            doc_type: Document type (auto-detected if None)
            extract_metadata: Whether to extract metadata

        Returns:
            Tuple of (content, metadata)
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Auto-detect document type if not provided
        if doc_type is None:
            extension = file_path.suffix.lower().lstrip('.')
            try:
                doc_type = DocumentType(extension)
            except ValueError:
                doc_type = DocumentType.TXT

        # Create base metadata
        metadata = DocumentMetadata(
            source=str(file_path),
            doc_type=doc_type
        )

        # Parse document
        parser = self.parsers.get(doc_type, self._parse_text)
        content = parser(file_path, metadata if extract_metadata else None)

        # Calculate provenance hash
        if extract_metadata:
            metadata.calculate_provenance(content)
            metadata.word_count = len(content.split())

        return content, metadata

    def _parse_pdf(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse PDF documents"""
        try:
            import PyPDF2
            content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                if metadata:
                    metadata.page_count = len(pdf_reader.pages)
                    if pdf_reader.metadata:
                        metadata.author = pdf_reader.metadata.get('/Author')
                        metadata.title = pdf_reader.metadata.get('/Title')

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"Page {page_num + 1}:\n{text}")

            return "\n\n".join(content)

        except ImportError:
            logger.warning("PyPDF2 not installed, trying pdfplumber")
            return self._parse_pdf_plumber(file_path, metadata)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return self._parse_pdf_plumber(file_path, metadata)

    def _parse_pdf_plumber(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse PDF with pdfplumber as fallback"""
        try:
            import pdfplumber
            content = []

            with pdfplumber.open(file_path) as pdf:
                if metadata:
                    metadata.page_count = len(pdf.pages)
                    if pdf.metadata:
                        metadata.author = pdf.metadata.get('Author')
                        metadata.title = pdf.metadata.get('Title')

                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content.append(f"Page {page_num + 1}:\n{text}")

                    # Extract tables if present
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = "\n".join(["\t".join(row) for row in table if row])
                        content.append(f"Table on Page {page_num + 1}:\n{table_text}")

            return "\n\n".join(content)

        except ImportError:
            logger.error("Neither PyPDF2 nor pdfplumber installed")
            raise ImportError("Install PyPDF2 or pdfplumber: pip install PyPDF2 pdfplumber")
        except Exception as e:
            logger.error(f"Error parsing PDF with pdfplumber: {e}")
            raise

    def _parse_html(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse HTML documents"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')

                # Extract metadata
                if metadata:
                    title_tag = soup.find('title')
                    if title_tag:
                        metadata.title = title_tag.text.strip()

                    meta_author = soup.find('meta', attrs={'name': 'author'})
                    if meta_author:
                        metadata.author = meta_author.get('content')

                    meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
                    if meta_keywords:
                        keywords = meta_keywords.get('content', '').split(',')
                        metadata.tags = [k.strip() for k in keywords]

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Extract text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)

                return text

        except ImportError:
            raise ImportError("BeautifulSoup not installed: pip install beautifulsoup4")
        except Exception as e:
            logger.error(f"Error parsing HTML: {e}")
            raise

    def _parse_docx(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse DOCX documents"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            content = []

            # Extract metadata
            if metadata:
                core_properties = doc.core_properties
                metadata.author = core_properties.author
                metadata.title = core_properties.title
                if core_properties.keywords:
                    metadata.tags = [k.strip() for k in core_properties.keywords.split(',')]

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells]
                    if any(row_content):
                        table_content.append('\t'.join(row_content))

                if table_content:
                    content.append(f"\nTable {table_idx + 1}:")
                    content.extend(table_content)

            return '\n\n'.join(content)

        except ImportError:
            raise ImportError("python-docx not installed: pip install python-docx")
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    def _parse_xlsx(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse XLSX spreadsheets"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"Sheet: {sheet_name}\n")

                # Extract data from sheet
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                        content.append(row_text)

                content.append("")  # Add blank line between sheets

            return '\n'.join(content)

        except ImportError:
            # Fallback to pandas if available
            try:
                import pandas as pd
                xlsx_data = pd.read_excel(file_path, sheet_name=None)

                content = []
                for sheet_name, df in xlsx_data.items():
                    content.append(f"Sheet: {sheet_name}\n")
                    content.append(df.to_string())
                    content.append("")

                return '\n'.join(content)
            except ImportError:
                raise ImportError("Install openpyxl or pandas: pip install openpyxl pandas")
        except Exception as e:
            logger.error(f"Error parsing XLSX: {e}")
            raise

    def _parse_csv(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse CSV files"""
        try:
            import csv

            content = []
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    content.append('\t'.join(row))

            return '\n'.join(content)

        except Exception as e:
            # Fallback to pandas if available
            try:
                import pandas as pd
                df = pd.read_csv(file_path)
                return df.to_string()
            except ImportError:
                logger.error(f"Error parsing CSV: {e}")
                raise

    def _parse_json(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)

            # Convert JSON to readable text format
            def json_to_text(obj, indent=0):
                text = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        text.append(' ' * indent + f"{key}:")
                        text.extend(json_to_text(value, indent + 2))
                elif isinstance(obj, list):
                    for item in obj:
                        text.append(' ' * indent + '-')
                        text.extend(json_to_text(item, indent + 2))
                else:
                    text.append(' ' * indent + str(obj))
                return text

            return '\n'.join(json_to_text(data))

        except Exception as e:
            logger.error(f"Error parsing JSON: {e}")
            raise

    def _parse_xml(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse XML files"""
        try:
            import xml.etree.ElementTree as ET

            tree = ET.parse(file_path)
            root = tree.getroot()

            def xml_to_text(element, indent=0):
                text = []
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                text.append(' ' * indent + f"{tag}:")

                if element.text and element.text.strip():
                    text.append(' ' * (indent + 2) + element.text.strip())

                for child in element:
                    text.extend(xml_to_text(child, indent + 2))

                return text

            return '\n'.join(xml_to_text(root))

        except Exception as e:
            logger.error(f"Error parsing XML: {e}")
            raise

    def _parse_text(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise

    def _parse_markdown(self, file_path: Path, metadata: Optional[DocumentMetadata]) -> str:
        """Parse Markdown files"""
        content = self._parse_text(file_path, metadata)

        # Extract metadata from front matter if present
        if metadata and content.startswith('---'):
            try:
                import yaml
                parts = content.split('---', 2)
                if len(parts) >= 3:
                    front_matter = yaml.safe_load(parts[1])
                    if isinstance(front_matter, dict):
                        metadata.title = front_matter.get('title')
                        metadata.author = front_matter.get('author')
                        metadata.tags = front_matter.get('tags', [])
                        metadata.custom_metadata.update(front_matter)
                    content = parts[2]
            except:
                pass

        return content


class DocumentProcessor:
    """Process and chunk documents for RAG system"""

    def __init__(
        self,
        strategy: ChunkingStrategy = ChunkingStrategy.SLIDING_WINDOW,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100,
        max_chunk_size: int = 2000
    ):
        """
        Initialize document processor

        Args:
            strategy: Chunking strategy to use
            chunk_size: Target chunk size in tokens
            chunk_overlap: Overlap between chunks in tokens
            min_chunk_size: Minimum chunk size (discard smaller)
            max_chunk_size: Maximum chunk size (split larger)
        """
        self.strategy = strategy
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.parser = DocumentParser()

        # For semantic chunking
        self.nlp = None
        self.tokenizer = None

    def process_file(
        self,
        file_path: Union[str, Path],
        doc_type: Optional[DocumentType] = None,
        custom_metadata: Optional[Dict] = None
    ) -> List[Document]:
        """
        Process a file into chunked documents

        Args:
            file_path: Path to the file
            doc_type: Document type (auto-detected if None)
            custom_metadata: Additional metadata to attach

        Returns:
            List of Document objects
        """
        # Parse document
        content, metadata = self.parser.parse(file_path, doc_type)

        # Add custom metadata
        if custom_metadata:
            metadata.custom_metadata.update(custom_metadata)

        # Chunk document
        return self.chunk_text(content, metadata)

    def chunk_text(
        self,
        text: str,
        metadata: DocumentMetadata
    ) -> List[Document]:
        """
        Chunk text based on selected strategy

        Args:
            text: Text to chunk
            metadata: Document metadata

        Returns:
            List of Document chunks
        """
        if self.strategy == ChunkingStrategy.FIXED_SIZE:
            chunks = self._fixed_size_chunking(text)
        elif self.strategy == ChunkingStrategy.SLIDING_WINDOW:
            chunks = self._sliding_window_chunking(text)
        elif self.strategy == ChunkingStrategy.SEMANTIC:
            chunks = self._semantic_chunking(text)
        elif self.strategy == ChunkingStrategy.SENTENCE_BASED:
            chunks = self._sentence_based_chunking(text)
        elif self.strategy == ChunkingStrategy.RECURSIVE_CHARACTER:
            chunks = self._recursive_character_chunking(text)
        elif self.strategy == ChunkingStrategy.PARAGRAPH_BASED:
            chunks = self._paragraph_based_chunking(text)
        elif self.strategy == ChunkingStrategy.TOKEN_BASED:
            chunks = self._token_based_chunking(text)
        else:
            raise ValueError(f"Unknown chunking strategy: {self.strategy}")

        # Create Document objects
        documents = []
        for i, (chunk_text, start_char, end_char) in enumerate(chunks):
            if len(chunk_text.split()) >= self.min_chunk_size // 5:  # Rough token estimate
                doc = Document(
                    content=chunk_text,
                    metadata=metadata,
                    chunk_id=i,
                    start_char=start_char,
                    end_char=end_char
                )
                documents.append(doc)

        logger.info(f"Created {len(documents)} chunks from document {metadata.source}")
        return documents

    def _fixed_size_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Fixed size chunking without overlap"""
        chunks = []
        words = text.split()

        for i in range(0, len(words), self.chunk_size):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)

            # Calculate character positions
            start_char = text.find(chunk_words[0]) if chunk_words else 0
            end_char = text.find(chunk_words[-1]) + len(chunk_words[-1]) if chunk_words else 0

            chunks.append((chunk_text, start_char, end_char))

        return chunks

    def _sliding_window_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Sliding window chunking with overlap"""
        chunks = []
        words = text.split()
        step = max(1, self.chunk_size - self.chunk_overlap)

        for i in range(0, len(words), step):
            chunk_words = words[i:i + self.chunk_size]
            if not chunk_words:
                break

            chunk_text = ' '.join(chunk_words)

            # Calculate character positions
            start_char = len(' '.join(words[:i])) + (1 if i > 0 else 0)
            end_char = start_char + len(chunk_text)

            chunks.append((chunk_text, start_char, end_char))

            # Stop if we've processed all words
            if i + self.chunk_size >= len(words):
                break

        return chunks

    def _semantic_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Semantic chunking based on sentence boundaries"""
        if self.nlp is None:
            try:
                import spacy
                self.nlp = spacy.load("en_core_web_sm")
            except:
                logger.warning("Spacy not available, using sentence-based chunking")
                return self._sentence_based_chunking(text)

        doc = self.nlp(text)
        chunks = []
        current_chunk = []
        current_size = 0
        current_start = 0

        for sent in doc.sents:
            sent_text = sent.text.strip()
            sent_size = len(sent_text.split())

            if current_size + sent_size > self.chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

                # Overlap: include last sentences in next chunk
                overlap_size = 0
                overlap_sents = []
                for s in reversed(current_chunk):
                    overlap_size += len(s.split())
                    overlap_sents.insert(0, s)
                    if overlap_size >= self.chunk_overlap:
                        break

                current_chunk = overlap_sents + [sent_text]
                current_size = overlap_size + sent_size
                current_start = sent.start_char
            else:
                current_chunk.append(sent_text)
                current_size += sent_size
                if not current_chunk:
                    current_start = sent.start_char

        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

        return chunks

    def _sentence_based_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Sentence-based chunking"""
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)

        chunks = []
        current_chunk = []
        current_size = 0
        current_start = 0

        for sent in sentences:
            sent_size = len(sent.split())

            if current_size + sent_size > self.chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

                # Add overlap
                overlap_size = 0
                overlap_sents = []
                for s in reversed(current_chunk):
                    overlap_size += len(s.split())
                    overlap_sents.insert(0, s)
                    if overlap_size >= self.chunk_overlap:
                        break

                current_chunk = overlap_sents + [sent]
                current_size = overlap_size + sent_size
                current_start = text.find(sent, current_start)
            else:
                if not current_chunk:
                    current_start = text.find(sent)
                current_chunk.append(sent)
                current_size += sent_size

        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

        return chunks

    def _recursive_character_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Recursive character text splitting"""
        separators = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""]

        def split_text(txt: str, start_pos: int = 0) -> List[Tuple[str, int, int]]:
            chunks = []

            # Find best separator
            separator = None
            for sep in separators:
                if sep in txt and len(txt) > self.chunk_size:
                    separator = sep
                    break

            if not separator or separator == "":
                # Base case: split by chunk size
                for i in range(0, len(txt), self.chunk_size):
                    chunk = txt[i:i + self.chunk_size]
                    chunks.append((chunk, start_pos + i, start_pos + i + len(chunk)))
                return chunks

            # Recursive case: split by separator
            parts = txt.split(separator)
            current_chunk = []
            current_size = 0
            chunk_start = start_pos

            for i, part in enumerate(parts):
                part_size = len(part)

                if current_size + part_size + len(separator) > self.chunk_size:
                    if current_chunk:
                        chunk_text = separator.join(current_chunk)
                        chunks.append((chunk_text, chunk_start, chunk_start + len(chunk_text)))

                    if part_size > self.chunk_size:
                        # Recursively split large parts
                        sub_chunks = split_text(part, chunk_start)
                        chunks.extend(sub_chunks)
                        chunk_start += sum(len(c[0]) for c in sub_chunks)
                    else:
                        current_chunk = [part]
                        current_size = part_size
                        chunk_start = start_pos + len(separator.join(parts[:i])) + len(separator) * i
                else:
                    current_chunk.append(part)
                    current_size += part_size + len(separator)

            if current_chunk:
                chunk_text = separator.join(current_chunk)
                chunks.append((chunk_text, chunk_start, chunk_start + len(chunk_text)))

            return chunks

        return split_text(text)

    def _paragraph_based_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Paragraph-based chunking"""
        # Split into paragraphs
        paragraphs = re.split(r'\n\s*\n', text)

        chunks = []
        current_chunk = []
        current_size = 0
        current_start = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_size = len(para.split())

            if current_size + para_size > self.chunk_size and current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

                # No overlap for paragraph-based chunking
                current_chunk = [para]
                current_size = para_size
                current_start = text.find(para, current_start + len(chunk_text))
            else:
                if not current_chunk:
                    current_start = text.find(para)
                current_chunk.append(para)
                current_size += para_size

        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            chunks.append((chunk_text, current_start, current_start + len(chunk_text)))

        return chunks

    def _token_based_chunking(self, text: str) -> List[Tuple[str, int, int]]:
        """Token-based chunking using a tokenizer"""
        if self.tokenizer is None:
            try:
                from transformers import AutoTokenizer
                self.tokenizer = AutoTokenizer.from_pretrained("bert-base-uncased")
            except ImportError:
                logger.warning("Transformers not available, using word-based chunking")
                return self._sliding_window_chunking(text)

        # Tokenize text
        tokens = self.tokenizer.tokenize(text)
        token_ids = self.tokenizer.convert_tokens_to_ids(tokens)

        chunks = []
        step = max(1, self.chunk_size - self.chunk_overlap)

        for i in range(0, len(token_ids), step):
            chunk_ids = token_ids[i:i + self.chunk_size]
            chunk_tokens = self.tokenizer.convert_ids_to_tokens(chunk_ids)
            chunk_text = self.tokenizer.convert_tokens_to_string(chunk_tokens)

            # Approximate character positions
            start_char = len(self.tokenizer.convert_tokens_to_string(tokens[:i]))
            end_char = start_char + len(chunk_text)

            chunks.append((chunk_text, start_char, end_char))

            if i + self.chunk_size >= len(token_ids):
                break

        return chunks

    def process_batch(
        self,
        file_paths: List[Union[str, Path]],
        doc_types: Optional[List[DocumentType]] = None,
        custom_metadata: Optional[Dict] = None
    ) -> List[Document]:
        """
        Process multiple files in batch

        Args:
            file_paths: List of file paths
            doc_types: List of document types (auto-detected if None)
            custom_metadata: Additional metadata to attach to all documents

        Returns:
            List of all Document chunks
        """
        all_documents = []

        if doc_types is None:
            doc_types = [None] * len(file_paths)

        for file_path, doc_type in zip(file_paths, doc_types):
            try:
                documents = self.process_file(file_path, doc_type, custom_metadata)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue

        logger.info(f"Processed {len(file_paths)} files into {len(all_documents)} chunks")
        return all_documents